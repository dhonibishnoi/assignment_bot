from docx import Document
import io

class DOCXService:
    @staticmethod
    async def generate_docx(assignment_text: str, topic: str) -> bytes:
        doc = Document()
        doc.add_heading(f'Assignment: {topic}', level=1)
        for para in assignment_text.split('\n\n'):
            doc.add_paragraph(para)
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.read()